from typing import List, Dict, Any, Optional
import pandas as pd
import numpy as np
from datetime import datetime
import jinja2
import pdfkit
import markdown
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
import base64

class ReportGenerator:
    def __init__(self, data: pd.DataFrame):
        self.data = data
        self.env = jinja2.Environment(
            loader=jinja2.FileSystemLoader('templates'),
            autoescape=jinja2.select_autoescape(['html', 'xml'])
        )

    def generate_report(
        self,
        title: str,
        description: str,
        sections: List[str],
        format: str = 'pdf',
        options: Dict[str, bool] = None
    ) -> BytesIO:
        """Generate a report in the specified format."""
        options = options or {
            'include_code': True,
            'include_charts': True,
            'include_statistics': True,
        }

        # Generate content for each section
        content = []
        for section in sections:
            section_content = self._generate_section_content(section, options)
            if section_content:
                content.append(section_content)

        # Prepare template context
        context = {
            'title': title,
            'description': description,
            'content': content,
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'options': options,
        }

        # Generate report in requested format
        if format == 'pdf':
            return self._generate_pdf(context)
        elif format == 'html':
            return self._generate_html(context)
        elif format == 'docx':
            return self._generate_docx(context)
        elif format == 'md':
            return self._generate_markdown(context)
        else:
            raise ValueError(f'Unsupported format: {format}')

    def _generate_section_content(
        self,
        section: str,
        options: Dict[str, bool]
    ) -> Optional[Dict[str, Any]]:
        """Generate content for a specific section."""
        if section == 'data_summary':
            return self._generate_data_summary(options)
        elif section == 'feature_analysis':
            return self._generate_feature_analysis(options)
        elif section == 'correlation_analysis':
            return self._generate_correlation_analysis(options)
        elif section == 'distribution_analysis':
            return self._generate_distribution_analysis(options)
        elif section == 'feature_importance':
            return self._generate_feature_importance(options)
        else:
            return None

    def _generate_data_summary(self, options: Dict[str, bool]) -> Dict[str, Any]:
        """Generate data summary section."""
        summary = {
            'title': 'Data Summary',
            'text': [],
            'tables': [],
            'charts': [],
            'code': [],
        }

        # Basic dataset information
        summary['text'].append(f"Dataset Shape: {self.data.shape[0]} rows × {self.data.shape[1]} columns")
        summary['text'].append(f"Memory Usage: {self.data.memory_usage(deep=True).sum() / 1024 / 1024:.2f} MB")

        if options['include_statistics']:
            # Generate descriptive statistics
            stats_df = self.data.describe(include='all')
            summary['tables'].append({
                'title': 'Descriptive Statistics',
                'data': stats_df.to_html(classes='table table-striped'),
            })

        if options['include_charts']:
            # Generate missing values chart
            missing_data = self.data.isnull().sum()
            if missing_data.any():
                plt.figure(figsize=(10, 6))
                missing_data.plot(kind='bar')
                plt.title('Missing Values by Column')
                plt.xticks(rotation=45)
                plt.tight_layout()
                
                # Save plot to base64
                buffer = BytesIO()
                plt.savefig(buffer, format='png')
                plt.close()
                buffer.seek(0)
                image = base64.b64encode(buffer.getvalue()).decode()
                summary['charts'].append({
                    'title': 'Missing Values Distribution',
                    'data': f'data:image/png;base64,{image}',
                })

        if options['include_code']:
            summary['code'].append({
                'title': 'Basic Data Analysis',
                'language': 'python',
                'code': '''
# Load and examine the dataset
df.info()
df.describe()

# Check missing values
df.isnull().sum()
''',
            })

        return summary

    def _generate_feature_analysis(self, options: Dict[str, bool]) -> Dict[str, Any]:
        """Generate feature analysis section."""
        analysis = {
            'title': 'Feature Analysis',
            'text': [],
            'tables': [],
            'charts': [],
            'code': [],
        }

        numeric_cols = self.data.select_dtypes(include=[np.number]).columns
        categorical_cols = self.data.select_dtypes(exclude=[np.number]).columns

        analysis['text'].append(f"Numeric Features: {len(numeric_cols)}")
        analysis['text'].append(f"Categorical Features: {len(categorical_cols)}")

        if options['include_statistics']:
            # Generate statistics for numeric features
            if len(numeric_cols) > 0:
                numeric_stats = self.data[numeric_cols].agg(['mean', 'std', 'min', 'max', 'skew'])
                analysis['tables'].append({
                    'title': 'Numeric Features Statistics',
                    'data': numeric_stats.to_html(classes='table table-striped'),
                })

            # Generate statistics for categorical features
            if len(categorical_cols) > 0:
                cat_stats = pd.DataFrame({
                    col: self.data[col].value_counts().head(5).to_dict()
                    for col in categorical_cols
                })
                analysis['tables'].append({
                    'title': 'Top 5 Categories per Feature',
                    'data': cat_stats.to_html(classes='table table-striped'),
                })

        if options['include_charts']:
            # Generate distribution plots for numeric features
            for col in numeric_cols[:5]:  # Limit to first 5 features
                plt.figure(figsize=(10, 6))
                sns.histplot(data=self.data, x=col, kde=True)
                plt.title(f'Distribution of {col}')
                plt.tight_layout()
                
                buffer = BytesIO()
                plt.savefig(buffer, format='png')
                plt.close()
                buffer.seek(0)
                image = base64.b64encode(buffer.getvalue()).decode()
                analysis['charts'].append({
                    'title': f'Distribution of {col}',
                    'data': f'data:image/png;base64,{image}',
                })

        if options['include_code']:
            analysis['code'].append({
                'title': 'Feature Analysis Code',
                'language': 'python',
                'code': '''
# Analyze numeric features
numeric_features = df.select_dtypes(include=[np.number])
numeric_stats = numeric_features.agg(['mean', 'std', 'min', 'max', 'skew'])

# Analyze categorical features
categorical_features = df.select_dtypes(exclude=[np.number])
for col in categorical_features:
    print(f"\\nValue counts for {col}:")
    print(df[col].value_counts().head())
''',
            })

        return analysis

    def _generate_correlation_analysis(self, options: Dict[str, bool]) -> Dict[str, Any]:
        """Generate correlation analysis section."""
        analysis = {
            'title': 'Correlation Analysis',
            'text': [],
            'tables': [],
            'charts': [],
            'code': [],
        }

        numeric_data = self.data.select_dtypes(include=[np.number])
        if len(numeric_data.columns) < 2:
            analysis['text'].append("Insufficient numeric features for correlation analysis")
            return analysis

        # Calculate correlations
        correlations = numeric_data.corr()

        if options['include_statistics']:
            # Find highest correlations
            high_correlations = []
            for i in range(len(correlations.columns)):
                for j in range(i+1, len(correlations.columns)):
                    corr = correlations.iloc[i, j]
                    if abs(corr) > 0.5:
                        high_correlations.append({
                            'feature1': correlations.columns[i],
                            'feature2': correlations.columns[j],
                            'correlation': corr,
                        })
            
            if high_correlations:
                high_corr_df = pd.DataFrame(high_correlations)
                analysis['tables'].append({
                    'title': 'High Correlations (|r| > 0.5)',
                    'data': high_corr_df.to_html(classes='table table-striped'),
                })

        if options['include_charts']:
            # Generate correlation heatmap
            plt.figure(figsize=(12, 8))
            sns.heatmap(correlations, annot=True, cmap='coolwarm', center=0)
            plt.title('Correlation Heatmap')
            plt.tight_layout()
            
            buffer = BytesIO()
            plt.savefig(buffer, format='png')
            plt.close()
            buffer.seek(0)
            image = base64.b64encode(buffer.getvalue()).decode()
            analysis['charts'].append({
                'title': 'Correlation Heatmap',
                'data': f'data:image/png;base64,{image}',
            })

        if options['include_code']:
            analysis['code'].append({
                'title': 'Correlation Analysis Code',
                'language': 'python',
                'code': '''
# Calculate correlations
correlations = df.select_dtypes(include=[np.number]).corr()

# Find high correlations
high_correlations = []
for i in range(len(correlations.columns)):
    for j in range(i+1, len(correlations.columns)):
        corr = correlations.iloc[i, j]
        if abs(corr) > 0.5:
            print(f"{correlations.columns[i]} vs {correlations.columns[j]}: {corr:.3f}")
''',
            })

        return analysis

    def _generate_distribution_analysis(self, options: Dict[str, bool]) -> Dict[str, Any]:
        """Generate distribution analysis section."""
        analysis = {
            'title': 'Distribution Analysis',
            'text': [],
            'tables': [],
            'charts': [],
            'code': [],
        }

        numeric_cols = self.data.select_dtypes(include=[np.number]).columns

        if options['include_statistics']:
            # Calculate distribution statistics
            dist_stats = pd.DataFrame({
                'skewness': self.data[numeric_cols].skew(),
                'kurtosis': self.data[numeric_cols].kurtosis(),
            })
            analysis['tables'].append({
                'title': 'Distribution Statistics',
                'data': dist_stats.to_html(classes='table table-striped'),
            })

        if options['include_charts']:
            # Generate box plots
            plt.figure(figsize=(12, 6))
            self.data[numeric_cols].boxplot()
            plt.title('Box Plots of Numeric Features')
            plt.xticks(rotation=45)
            plt.tight_layout()
            
            buffer = BytesIO()
            plt.savefig(buffer, format='png')
            plt.close()
            buffer.seek(0)
            image = base64.b64encode(buffer.getvalue()).decode()
            analysis['charts'].append({
                'title': 'Box Plots',
                'data': f'data:image/png;base64,{image}',
            })

            # Generate Q-Q plots for first few numeric features
            for col in numeric_cols[:3]:
                plt.figure(figsize=(8, 6))
                stats.probplot(self.data[col].dropna(), dist="norm", plot=plt)
                plt.title(f'Q-Q Plot of {col}')
                plt.tight_layout()
                
                buffer = BytesIO()
                plt.savefig(buffer, format='png')
                plt.close()
                buffer.seek(0)
                image = base64.b64encode(buffer.getvalue()).decode()
                analysis['charts'].append({
                    'title': f'Q-Q Plot of {col}',
                    'data': f'data:image/png;base64,{image}',
                })

        if options['include_code']:
            analysis['code'].append({
                'title': 'Distribution Analysis Code',
                'language': 'python',
                'code': '''
# Calculate distribution statistics
numeric_features = df.select_dtypes(include=[np.number])
dist_stats = pd.DataFrame({
    'skewness': numeric_features.skew(),
    'kurtosis': numeric_features.kurtosis(),
})

# Generate distribution plots
import scipy.stats as stats
for column in numeric_features.columns:
    stats.probplot(df[column].dropna(), dist="norm", plot=plt)
    plt.title(f'Q-Q Plot of {column}')
    plt.show()
''',
            })

        return analysis

    def _generate_feature_importance(self, options: Dict[str, bool]) -> Dict[str, Any]:
        """Generate feature importance section using a random forest model."""
        analysis = {
            'title': 'Feature Importance Analysis',
            'text': [],
            'tables': [],
            'charts': [],
            'code': [],
        }

        numeric_data = self.data.select_dtypes(include=[np.number])
        if len(numeric_data.columns) < 2:
            analysis['text'].append("Insufficient numeric features for importance analysis")
            return analysis

        try:
            from sklearn.ensemble import RandomForestRegressor
            from sklearn.preprocessing import StandardScaler

            # Prepare data
            X = numeric_data.fillna(numeric_data.mean())
            target = X.columns[-1]  # Use last column as target for demonstration
            y = X[target]
            X = X.drop(columns=[target])

            # Scale features
            scaler = StandardScaler()
            X_scaled = scaler.fit_transform(X)
            X_scaled = pd.DataFrame(X_scaled, columns=X.columns)

            # Train model
            model = RandomForestRegressor(n_estimators=100, random_state=42)
            model.fit(X_scaled, y)

            # Get feature importance
            importance_df = pd.DataFrame({
                'feature': X.columns,
                'importance': model.feature_importances_,
            }).sort_values('importance', ascending=False)

            if options['include_statistics']:
                analysis['tables'].append({
                    'title': 'Feature Importance Scores',
                    'data': importance_df.to_html(classes='table table-striped'),
                })

            if options['include_charts']:
                # Generate feature importance plot
                plt.figure(figsize=(10, 6))
                sns.barplot(data=importance_df, x='importance', y='feature')
                plt.title('Feature Importance (Random Forest)')
                plt.tight_layout()
                
                buffer = BytesIO()
                plt.savefig(buffer, format='png')
                plt.close()
                buffer.seek(0)
                image = base64.b64encode(buffer.getvalue()).decode()
                analysis['charts'].append({
                    'title': 'Feature Importance Plot',
                    'data': f'data:image/png;base64,{image}',
                })

            if options['include_code']:
                analysis['code'].append({
                    'title': 'Feature Importance Analysis Code',
                    'language': 'python',
                    'code': '''
# Prepare data
X = df.select_dtypes(include=[np.number]).fillna(df.mean())
target = X.columns[-1]  # Use last column as target
y = X[target]
X = X.drop(columns=[target])

# Scale features
from sklearn.preprocessing import StandardScaler
scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)

# Train random forest
from sklearn.ensemble import RandomForestRegressor
model = RandomForestRegressor(n_estimators=100, random_state=42)
model.fit(X_scaled, y)

# Get feature importance
importance_df = pd.DataFrame({
    'feature': X.columns,
    'importance': model.feature_importances_,
}).sort_values('importance', ascending=False)
''',
                })

        except Exception as e:
            analysis['text'].append(f"Error in feature importance analysis: {str(e)}")

        return analysis

    def _generate_pdf(self, context: Dict[str, Any]) -> BytesIO:
        """Generate PDF report."""
        template = self.env.get_template('report_template.html')
        html = template.render(**context)
        
        output = BytesIO()
        pdfkit.from_string(html, output)
        output.seek(0)
        return output

    def _generate_html(self, context: Dict[str, Any]) -> BytesIO:
        """Generate HTML report."""
        template = self.env.get_template('report_template.html')
        html = template.render(**context)
        
        output = BytesIO()
        output.write(html.encode('utf-8'))
        output.seek(0)
        return output

    def _generate_docx(self, context: Dict[str, Any]) -> BytesIO:
        """Generate Word document report."""
        doc = Document()
        
        # Add title
        doc.add_heading(context['title'], 0)
        doc.add_paragraph(context['description'])
        doc.add_paragraph(f"Generated at: {context['generated_at']}")
        
        # Add content sections
        for section in context['content']:
            doc.add_heading(section['title'], 1)
            
            # Add text
            for text in section['text']:
                doc.add_paragraph(text)
            
            # Add tables
            for table in section['tables']:
                doc.add_heading(table['title'], 2)
                # Convert HTML table to docx table (simplified)
                doc.add_paragraph(table['data'])  # This needs proper HTML to DOCX conversion
            
            # Add charts
            for chart in section['charts']:
                doc.add_heading(chart['title'], 2)
                # Convert base64 image to docx image
                image_data = base64.b64decode(chart['data'].split(',')[1])
                image_stream = BytesIO(image_data)
                doc.add_picture(image_stream, width=Inches(6))
            
            # Add code
            for code_block in section['code']:
                doc.add_heading(code_block['title'], 2)
                doc.add_paragraph(code_block['code'], style='Code')
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    def _generate_markdown(self, context: Dict[str, Any]) -> BytesIO:
        """Generate Markdown report."""
        md_content = [
            f"# {context['title']}",
            context['description'],
            f"Generated at: {context['generated_at']}",
            "",
        ]
        
        for section in context['content']:
            md_content.extend([
                f"## {section['title']}",
                "",
            ])
            
            # Add text
            for text in section['text']:
                md_content.extend([text, ""])
            
            # Add tables
            for table in section['tables']:
                md_content.extend([
                    f"### {table['title']}",
                    table['data'],  # This needs proper HTML to MD table conversion
                    "",
                ])
            
            # Add charts
            for chart in section['charts']:
                md_content.extend([
                    f"### {chart['title']}",
                    f"![{chart['title']}]({chart['data']})",
                    "",
                ])
            
            # Add code
            for code_block in section['code']:
                md_content.extend([
                    f"### {code_block['title']}",
                    "```" + code_block['language'],
                    code_block['code'],
                    "```",
                    "",
                ])
        
        output = BytesIO()
        output.write('\n'.join(md_content).encode('utf-8'))
        output.seek(0)
        return output 